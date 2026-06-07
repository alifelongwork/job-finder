#!/usr/bin/env python3
"""jobsdb.py — local job-pipeline database CLI.

Single source of truth for the job list. See database.md for the full contract.
Core commands (Step 1): init, candidate add/list/show, category set,
upsert-batch, query, stats.  (reverify / mark arrive in Step 2; export in Step 5.)

Dependency-free: standard library only (sqlite3, argparse, json, csv, ...).
"""

import argparse
import datetime
import hashlib
import json
import os
import re
import sqlite3
import sys

# Make stdout UTF-8 so unicode (★, accented names, etc.) prints on Windows consoles.
try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

HERE = os.path.dirname(os.path.abspath(__file__))
# DB location: JOBSDB_PATH env var overrides the default (enables testing against a
# throwaway DB and lets multiple candidates keep separate databases).
DB_PATH = os.environ.get("JOBSDB_PATH", os.path.join(HERE, "jobs.db"))
SCHEMA_PATH = os.path.join(HERE, "schema.sql")

CANDIDATE_FIELDS = {
    "name", "email", "location_constraint", "citizenship", "clearance",
    "comp_floor", "comp_target", "resume_path", "notes",
}
INT_FIELDS = {"comp_floor", "comp_target"}
VALID_TAGS = {"verified", "wrong_location", "aggregator", "unverified"}
VALID_STATUS = {"new", "active", "applied", "expired", "rejected", "ignored"}
TERMINAL_STATUS = {"applied", "ignored", "rejected"}
# Dead / set-aside statuses hidden from default query+export views (opt in with --all).
# Kept in the DB (history + dedup), just not shown unless asked.
INACTIVE_STATUS = {"expired", "rejected", "ignored"}
# Company-level verification (the analog of a job's verification_tag), set via
# `company verify`. feed_verified = a live ATS JSON feed resolved; careers_only = a
# careers page exists but no clean feed; unresolved = couldn't resolve (transient/unknown
# — recheck later); unverified = explicitly checked and no hiring surface found.
COMPANY_VERIFY_STATUS = {"feed_verified", "careers_only", "unresolved", "unverified"}
# Columns added to `companies` after the original schema shipped. _migrate() adds any that
# are missing, so an existing jobs.db upgrades in place (init won't — it skips populated DBs).
COMPANY_MIGRATIONS = [
    ("verification_status", "TEXT"),
    ("last_verified", "TEXT"),
    ("open_roles", "INTEGER"),
]


# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------
def today():
    return datetime.date.today().isoformat()


def days_since(iso):
    """Whole days from an ISO date to today, or None if missing/unparseable."""
    if not iso:
        return None
    try:
        return (datetime.date.today() - datetime.date.fromisoformat(iso)).days
    except ValueError:
        return None


def verified_age(iso):
    """Short label for how long ago a job was last verified: never / today / Nd."""
    if not iso:
        return "never"
    d = days_since(iso)
    if d is None:
        return "?"
    return "today" if d <= 0 else f"{d}d"


def slugify(text):
    s = re.sub(r"[^a-z0-9]+", "_", (text or "").lower()).strip("_")
    return s or "candidate"


def parse_comp(raw):
    """Parse an annual-USD comp value into an int, honoring a k/M suffix.

    '120000' -> 120000, '$120,000' -> 120000, '120k' -> 120000, '1.2M' -> 1200000.
    Returns None when no number is present. (A plain integer like '80000' is taken
    literally; the suffix is what scales it, so '80k' is NOT silently read as 80.)
    """
    s = (raw or "").strip().lower().replace(",", "").replace("$", "").replace("_", "")
    m = re.search(r"(\d+(?:\.\d+)?)\s*([km]?)", s)
    if not m:
        return None
    num = float(m.group(1))
    if m.group(2) == "k":
        num *= 1_000
    elif m.group(2) == "m":
        num *= 1_000_000
    return int(round(num))


def _migrate(conn):
    """Idempotently add post-1.0 columns to an existing DB. Runs on every connect() (the
    universal chokepoint) because `init` short-circuits on a populated DB and so never sees
    it. Steady state is one cheap PRAGMA + set-diff and no write; each ALTER fires once per
    DB. Column names are hard-coded constants (COMPANY_MIGRATIONS) — never user input."""
    cols = {r["name"] for r in conn.execute("PRAGMA table_info(companies)")}
    added = False
    for name, decl in COMPANY_MIGRATIONS:
        if name not in cols:
            conn.execute(f"ALTER TABLE companies ADD COLUMN {name} {decl}")
            added = True
    if added:
        conn.commit()


def connect():
    if not os.path.exists(DB_PATH):
        sys.exit("Database not found. Run:  python jobsdb.py init")
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON")
    _migrate(conn)
    return conn


def get_candidate(conn, slug):
    row = conn.execute("SELECT * FROM candidates WHERE slug = ?", (slug,)).fetchone()
    if not row:
        sys.exit(f"No candidate with slug '{slug}'. See:  python jobsdb.py candidate list")
    return row


def make_dedup_key(job):
    """Return the job's dedup_key, computing the aggregator fallback if absent."""
    key = job.get("dedup_key")
    if key:
        return key.lower()
    plat, slug, jid = job.get("ats_platform"), job.get("ats_slug"), job.get("ats_job_id")
    if plat and slug and jid:
        return f"{plat}:{slug}:{jid}".lower()
    seed = "|".join([job.get("company", ""), job.get("title", ""), job.get("location", "")])
    return "agg:" + hashlib.sha1(seed.encode("utf-8")).hexdigest()[:12]


# ---------------------------------------------------------------------------
# init
# ---------------------------------------------------------------------------
def cmd_init(args):
    if os.path.exists(DB_PATH) and not args.force:
        # Allow re-running init to add missing tables, but never silently wipe data.
        existing = sqlite3.connect(DB_PATH).execute(
            "SELECT count(*) FROM sqlite_master WHERE type='table'"
        ).fetchone()[0]
        if existing:
            print(f"Database already exists at {DB_PATH} ({existing} tables). "
                  "Schema is idempotent; nothing to do. Use --force to recreate from scratch.")
            return
    if args.force and os.path.exists(DB_PATH):
        os.remove(DB_PATH)
        print("Removed existing database (--force).")
    with open(SCHEMA_PATH, encoding="utf-8") as f:
        schema = f.read()
    conn = sqlite3.connect(DB_PATH)
    conn.executescript(schema)
    conn.commit()
    conn.close()
    print(f"Initialized database at {DB_PATH}")


# ---------------------------------------------------------------------------
# candidate
# ---------------------------------------------------------------------------
def parse_fields(field_args):
    fields = {}
    for item in field_args or []:
        if "=" not in item:
            sys.exit(f"--field must be key=value, got: {item}")
        k, v = item.split("=", 1)
        k = k.strip()
        if k not in CANDIDATE_FIELDS:
            sys.exit(f"Unknown candidate field '{k}'. Allowed: {sorted(CANDIDATE_FIELDS)}")
        v = v.strip()
        if k in INT_FIELDS:
            v = parse_comp(v)
        fields[k] = v
    return fields


def cmd_candidate_add(args):
    fields = parse_fields(args.field)
    if args.resume:
        fields["resume_path"] = args.resume
    name = fields.get("name")
    slug = args.slug or slugify(name) if (args.slug or name) else None
    if not slug:
        sys.exit("Provide --slug or --field name=... to identify the candidate.")

    conn = connect()
    existing = conn.execute("SELECT * FROM candidates WHERE slug = ?", (slug,)).fetchone()
    now = today()
    if existing:
        if fields:
            sets = ", ".join(f"{k} = ?" for k in fields)
            conn.execute(
                f"UPDATE candidates SET {sets}, updated_at = ? WHERE slug = ?",
                [*fields.values(), now, slug],
            )
            conn.commit()
        print(f"Updated candidate '{slug}' ({len(fields)} field(s)).")
    else:
        if "name" not in fields:
            fields["name"] = slug.replace("_", " ").title()
        cols = ["slug", *fields.keys(), "created_at", "updated_at"]
        vals = [slug, *fields.values(), now, now]
        conn.execute(
            f"INSERT INTO candidates ({', '.join(cols)}) VALUES ({', '.join('?' * len(vals))})",
            vals,
        )
        conn.commit()
        print(f"Added candidate '{slug}'.")
    conn.close()


def cmd_candidate_list(args):
    conn = connect()
    rows = conn.execute(
        "SELECT slug, name, location_constraint, "
        "(SELECT count(*) FROM jobs j WHERE j.candidate_id = c.id) AS jobs "
        "FROM candidates c ORDER BY name"
    ).fetchall()
    conn.close()
    if not rows:
        print("No candidates yet.  python jobsdb.py candidate add --resume <path> --field name=...")
        return
    for r in rows:
        print(f"  {r['slug']:<18} {r['name']:<22} jobs={r['jobs']:<4} "
              f"{r['location_constraint'] or ''}")


def cmd_candidate_show(args):
    conn = connect()
    c = get_candidate(conn, args.slug)
    print(f"# {c['name']}  ({c['slug']})")
    for k in ("email", "location_constraint", "citizenship", "clearance",
              "comp_floor", "comp_target", "resume_path", "notes"):
        if c[k] not in (None, ""):
            print(f"  {k:<20} {c[k]}")
    cats = conn.execute(
        "SELECT rank, label, keywords FROM candidate_categories "
        "WHERE candidate_id = ? ORDER BY rank", (c["id"],),
    ).fetchall()
    if cats:
        print("  categories:")
        for cat in cats:
            print(f"    {cat['rank']}. {cat['label']}"
                  + (f"  [{cat['keywords']}]" if cat["keywords"] else ""))
    conn.close()


# ---------------------------------------------------------------------------
# category set
# ---------------------------------------------------------------------------
def cmd_category_set(args):
    with open(args.json, encoding="utf-8") as f:
        cats = json.load(f)
    if not isinstance(cats, list):
        sys.exit("categories JSON must be a list of {rank, label, keywords}.")
    conn = connect()
    c = get_candidate(conn, args.candidate)
    conn.execute("DELETE FROM candidate_categories WHERE candidate_id = ?", (c["id"],))
    for cat in cats:
        conn.execute(
            "INSERT INTO candidate_categories (candidate_id, rank, label, keywords) "
            "VALUES (?, ?, ?, ?)",
            (c["id"], cat["rank"], cat["label"], cat.get("keywords")),
        )
    conn.commit()
    conn.close()
    print(f"Set {len(cats)} categories for '{args.candidate}'.")


# ---------------------------------------------------------------------------
# upsert-batch
# ---------------------------------------------------------------------------
def upsert_company(conn, job):
    name = job.get("company")
    if not name:
        return None
    row = conn.execute("SELECT id FROM companies WHERE name = ?", (name,)).fetchone()
    # Text fields update only when a non-empty value is supplied (never clobber to blank).
    text_fields = {
        "careers_url": job.get("careers_url"),
        "ats_platform": job.get("ats_platform"),
        "ats_slug": job.get("ats_slug"),
        "warm_path": job.get("warm_path"),  # referral/contact note (drives Tier-1 sequencing)
    }
    multi_region = 1 if job.get("multi_region") else 0
    if row:
        # Only fill blanks / update non-null incoming text values.
        sets = {k: v for k, v in text_fields.items() if v not in (None, "")}
        # multi_region is a sticky flag: only ever flip it ON — a later scan that omits
        # the flag must not silently clear a company already marked multi-region.
        if multi_region:
            sets["multi_region"] = 1
        if sets:
            conn.execute(
                f"UPDATE companies SET {', '.join(f'{k}=?' for k in sets)} WHERE id=?",
                [*sets.values(), row["id"]],
            )
        return row["id"]
    cur = conn.execute(
        "INSERT INTO companies (name, careers_url, ats_platform, ats_slug, multi_region, "
        "warm_path) VALUES (?, ?, ?, ?, ?, ?)",
        (name, text_fields["careers_url"], text_fields["ats_platform"],
         text_fields["ats_slug"], multi_region, text_fields["warm_path"]),
    )
    return cur.lastrowid


def replace_contacts(conn, job_id, contacts):
    conn.execute("DELETE FROM contacts WHERE job_id = ?", (job_id,))
    for ct in contacts or []:
        conn.execute(
            "INSERT INTO contacts (job_id, name, title, priority, contact_type, hook, "
            "action, confirmed, notes) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (job_id, ct.get("name"), ct.get("title"), ct.get("priority"),
             ct.get("contact_type"), ct.get("hook"), ct.get("action"),
             1 if ct.get("confirmed") else 0, ct.get("notes")),
        )


def cmd_upsert_batch(args):
    with open(args.file, encoding="utf-8") as f:
        batch = json.load(f)
    run_date = batch.get("run_date") or today()
    conn = connect()
    c = get_candidate(conn, batch["candidate"])
    cid = c["id"]

    n_new = n_upd = 0
    for job in batch.get("jobs", []):
        tag = job.get("verification_tag")
        if tag not in VALID_TAGS:
            sys.exit(f"Job '{job.get('title')}' has invalid verification_tag '{tag}'.")
        if not job.get("title"):
            sys.exit(f"Job for company '{job.get('company')}' is missing required "
                     "field 'title' (NOT NULL).")
        key = make_dedup_key(job)
        company_id = upsert_company(conn, job)
        # 'verified' and 'wrong_location' were both confirmed live on the company
        # surface (location was read from it) -> they set last_verified. 'aggregator'
        # and 'unverified' were NOT confirmed live, so they stay due for re-verification.
        verified_now = run_date if tag in ("verified", "wrong_location") else None

        existing = conn.execute(
            "SELECT id, status, last_seen, last_verified "
            "FROM jobs WHERE candidate_id = ? AND dedup_key = ?",
            (cid, key),
        ).fetchone()

        common = {
            "company_id": company_id,
            "title": job.get("title"),
            "url": job.get("url"),
            "ats_platform": job.get("ats_platform"),
            "ats_job_id": job.get("ats_job_id"),
            "location": job.get("location"),
            "remote_type": job.get("remote_type"),
            "location_match": 1 if job.get("location_match") else 0,
            "comp_min": job.get("comp_min"),
            "comp_max": job.get("comp_max"),
            "posting_date": job.get("posting_date"),
            "verification_tag": tag,
            "tier": job.get("tier"),
            "category_label": job.get("category_label"),
            "fit_summary": job.get("fit_summary"),
            "screening_risks": job.get("screening_risks"),
        }

        if existing:
            updates = dict(common)
            # last_seen / last_verified only ever move forward (guard against
            # replaying an older scan batch with a back-dated run_date).
            prev_seen = existing["last_seen"]
            updates["last_seen"] = max(prev_seen, run_date) if prev_seen else run_date
            if verified_now:
                prev_ver = existing["last_verified"]
                updates["last_verified"] = max(prev_ver, verified_now) if prev_ver else verified_now
            # Preserve terminal status; revive a reappearing expired job.
            if existing["status"] not in TERMINAL_STATUS:
                if existing["status"] == "expired" and tag == "verified":
                    updates["status"] = "active"
                elif existing["status"] == "new":
                    updates["status"] = "active"
            sets = ", ".join(f"{k}=?" for k in updates)
            conn.execute(f"UPDATE jobs SET {sets} WHERE id=?",
                         [*updates.values(), existing["id"]])
            job_id = existing["id"]
            n_upd += 1
        else:
            cols = {
                "candidate_id": cid, "dedup_key": key, "status": "new",
                "first_seen": run_date, "last_seen": run_date,
                "last_verified": verified_now, **common,
            }
            conn.execute(
                f"INSERT INTO jobs ({', '.join(cols)}) "
                f"VALUES ({', '.join('?' * len(cols))})",
                list(cols.values()),
            )
            job_id = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
            n_new += 1

        if "contacts" in job:
            replace_contacts(conn, job_id, job["contacts"])

    found = n_new + n_upd
    conn.execute(
        "INSERT INTO search_runs (candidate_id, run_date, num_found, num_new, "
        "num_updated, num_expired, notes) VALUES (?, ?, ?, ?, ?, ?, ?)",
        (cid, run_date, found, n_new, n_upd, 0, batch.get("notes")),
    )
    conn.commit()
    conn.close()
    summary = {"found": found, "new": n_new, "updated": n_upd}
    print(json.dumps(summary))
    print(f"Upserted {found} job(s) for '{batch['candidate']}': "
          f"{n_new} new, {n_upd} updated.")


# ---------------------------------------------------------------------------
# query
# ---------------------------------------------------------------------------
QUERY_ORDER = """
ORDER BY
  CASE WHEN j.tier IS NULL THEN 9 ELSE j.tier END ASC,
  CASE j.verification_tag WHEN 'verified' THEN 0 WHEN 'aggregator' THEN 1
       WHEN 'unverified' THEN 2 WHEN 'wrong_location' THEN 3 ELSE 4 END ASC,
  CASE WHEN j.posting_date IS NULL OR j.posting_date = '' THEN 1 ELSE 0 END ASC,
  j.posting_date ASC
"""

TAG_SHORT = {"verified": "OK", "wrong_location": "WRONG-LOC",
             "aggregator": "AGG", "unverified": "UNV"}


def job_filter_clause(conn, args):
    """Shared WHERE builder for query/export (filters keyed on `j.` alias)."""
    where, params = [], []
    if getattr(args, "candidate", None):
        where.append("j.candidate_id = ?"); params.append(get_candidate(conn, args.candidate)["id"])
    if getattr(args, "category", None):
        where.append("j.category_label LIKE ?"); params.append(f"%{args.category}%")
    if getattr(args, "tier", None):
        where.append("j.tier = ?"); params.append(args.tier)
    if getattr(args, "status", None):
        where.append("j.status = ?"); params.append(args.status)
    elif not getattr(args, "all", False):
        # No explicit status and no --all: show only the live/actionable pipeline.
        ph = ", ".join("?" * len(INACTIVE_STATUS))
        where.append(f"j.status NOT IN ({ph})")
        params.extend(sorted(INACTIVE_STATUS))
    if getattr(args, "verification", None):
        where.append("j.verification_tag = ?"); params.append(args.verification)
    if getattr(args, "location_match", None):
        where.append("j.location_match = ?")
        params.append(1 if args.location_match == "yes" else 0)
    if getattr(args, "since", None):
        where.append("j.last_seen >= ?"); params.append(args.since)
    return where, params


def cmd_query(args):
    conn = connect()
    where, params = job_filter_clause(conn, args)

    sql = (
        "SELECT j.id, j.tier, j.status, j.verification_tag, j.location_match, "
        "COALESCE(co.name, '') AS company, j.title, j.location, j.posting_date, "
        "j.comp_min, j.comp_max, j.url, j.category_label, j.last_verified "
        "FROM jobs j LEFT JOIN companies co ON co.id = j.company_id"
    )
    if where:
        sql += " WHERE " + " AND ".join(where)
    sql += QUERY_ORDER
    if args.limit:
        sql += f" LIMIT {int(args.limit)}"
    rows = conn.execute(sql, params).fetchall()
    conn.close()

    if args.format == "json":
        print(json.dumps([dict(r) for r in rows], indent=2))
        return
    if not rows:
        print("(no matching jobs)")
        return

    def comp(r):
        lo, hi = r["comp_min"], r["comp_max"]
        if lo and hi:
            return f"{lo // 1000}-{hi // 1000}k"
        if lo:
            return f"{lo // 1000}k+"
        return ""

    print(f"{'id':>4}  {'T':<1} {'tag':<9} {'status':<8} {'company':<22} "
          f"{'title':<34} {'location':<18} {'posted':<10} {'verif':<7} comp")
    print("-" * 140)
    for r in rows:
        tier = r["tier"] if r["tier"] else "-"
        d = days_since(r["last_verified"])
        # Flag rows that are stale-ish or never verified so freshness is visible at a glance.
        stale = r["last_verified"] is None or (d is not None and d > 7)
        verif = verified_age(r["last_verified"]) + ("!" if stale else "")
        print(f"{r['id']:>4}  {tier:<1} {TAG_SHORT.get(r['verification_tag'], '?'):<9} "
              f"{r['status']:<8} {(r['company'] or '')[:22]:<22} "
              f"{(r['title'] or '')[:34]:<34} {(r['location'] or '')[:18]:<18} "
              f"{(r['posting_date'] or '?'):<10} {verif:<7} {comp(r)}")
    hidden = "" if (args.status or args.all) else "  ·  expired/rejected/ignored hidden (--all to show)"
    print(f"\n{len(rows)} job(s).   (verif = age since last live-check; ! = never or >7d — re-verify before applying){hidden}")


# ---------------------------------------------------------------------------
# stats
# ---------------------------------------------------------------------------
def cmd_stats(args):
    conn = connect()
    c = get_candidate(conn, args.candidate)
    cid = c["id"]

    def counts(col):
        return conn.execute(
            f"SELECT {col} AS k, count(*) AS n FROM jobs WHERE candidate_id = ? "
            f"GROUP BY {col} ORDER BY n DESC", (cid,),
        ).fetchall()

    total = conn.execute("SELECT count(*) FROM jobs WHERE candidate_id = ?",
                         (cid,)).fetchone()[0]
    print(f"# Pipeline for {c['name']} ({c['slug']}) — {total} job(s)\n")

    print("by tier:")
    for r in counts("tier"):
        print(f"  Tier {r['k'] if r['k'] else '-'}: {r['n']}")
    print("by status:")
    for r in counts("status"):
        print(f"  {r['k']:<10} {r['n']}")
    print("by verification:")
    for r in counts("verification_tag"):
        print(f"  {r['k']:<14} {r['n']}")
    matched = conn.execute(
        "SELECT count(*) FROM jobs WHERE candidate_id=? AND location_match=1", (cid,)
    ).fetchone()[0]
    print(f"location match: {matched}/{total} passed")
    conn.close()


# ---------------------------------------------------------------------------
# reverify list
# ---------------------------------------------------------------------------
def cmd_reverify(args):
    """Emit live (new/active) jobs whose verification is stale, for the agent to re-check.

    The CLI does no network I/O. It tells the agent which URLs to re-fetch; the agent
    fetches each and records the outcome with `mark` (--verified for live, --status
    expired for dead).
    """
    if args.action != "list":
        sys.exit("Usage: jobsdb.py reverify list --candidate <slug> [--stale-days N]")
    conn = connect()
    c = get_candidate(conn, args.candidate)
    today_str = today()
    cutoff = (datetime.date.today()
              - datetime.timedelta(days=args.stale_days)).isoformat()
    # A job is due for re-verification if: never verified, OR older than the window,
    # OR it's Tier 1/2 and hasn't been re-checked yet today (those are the roles the
    # candidate would actually act on, so we keep them as fresh as a sweep allows).
    rows = conn.execute(
        "SELECT j.id, j.title, j.url, j.last_verified, j.status, j.tier, "
        "COALESCE(co.name,'') AS company "
        "FROM jobs j LEFT JOIN companies co ON co.id = j.company_id "
        "WHERE j.candidate_id = ? AND j.status IN ('new','active') "
        "AND (j.last_verified IS NULL OR j.last_verified < ? "
        "     OR (j.tier IN (1, 2) AND j.last_verified < ?)) "
        "ORDER BY CASE WHEN j.tier IS NULL THEN 9 ELSE j.tier END, j.last_verified",
        (c["id"], cutoff, today_str),
    ).fetchall()
    conn.close()

    if args.format == "json":
        print(json.dumps([dict(r) for r in rows], indent=2))
        return
    if not rows:
        print(f"No jobs need re-verification (stale-days={args.stale_days}).")
        return
    print(f"{len(rows)} job(s) need re-verification "
          f"(verified before {cutoff}, or Tier 1/2 not yet re-checked today, or never):\n")
    for r in rows:
        tier = f"T{r['tier']}" if r["tier"] else "T-"
        print(f"  id={r['id']:<4} {tier} {(r['company'] or '')[:20]:<20} "
              f"{(r['title'] or '')[:34]:<34} last_verified={r['last_verified'] or 'never'}")
        print(f"        {r['url']}")
    print("\nRe-fetch each URL, then:  jobsdb.py mark <id> --verified   (live)")
    print("                          jobsdb.py mark <id> --status expired   (dead)")


# ---------------------------------------------------------------------------
# mark
# ---------------------------------------------------------------------------
def cmd_mark(args):
    conn = connect()
    job = conn.execute("SELECT * FROM jobs WHERE id = ?", (args.job_id,)).fetchone()
    if not job:
        sys.exit(f"No job with id {args.job_id}.")

    updates = {}
    if args.status:
        updates["status"] = args.status
        if args.status == "applied" and not args.applied_date:
            updates["applied_date"] = today()
    if args.applied_date:
        updates["applied_date"] = args.applied_date
    if args.verified:
        updates["last_verified"] = today()
        # A confirmed-live job that wasn't terminal becomes/stays active.
        if job["status"] in ("new", "expired") and not args.status:
            updates["status"] = "active"
    if args.resume:
        updates["resume_path"] = args.resume
    if args.cover:
        updates["cover_letter_path"] = args.cover
    if args.note:
        existing = (job["notes"] + "\n") if job["notes"] else ""
        updates["notes"] = existing + f"[{today()}] {args.note}"

    if not updates:
        sys.exit("Nothing to update. Provide --status/--verified/--resume/--cover/"
                 "--applied-date/--note.")

    sets = ", ".join(f"{k}=?" for k in updates)
    conn.execute(f"UPDATE jobs SET {sets} WHERE id=?", [*updates.values(), args.job_id])
    conn.commit()
    conn.close()
    print(f"Updated job {args.job_id}: " + ", ".join(f"{k}={v}" for k, v in updates.items()))


# ---------------------------------------------------------------------------
# export  (CSV + Markdown = stdlib; .docx = lazy python-docx)
# ---------------------------------------------------------------------------
EXPORT_COLS = [
    "id", "tier", "status", "verification_tag", "company", "title", "location",
    "remote_type", "location_match", "comp_min", "comp_max", "posting_date",
    "category_label", "url", "first_seen", "last_seen", "last_verified",
    "applied_date", "fit_summary", "screening_risks", "resume_path", "cover_letter_path",
]


def _export_rows(conn, args):
    where, params = job_filter_clause(conn, args)
    sql = ("SELECT j.*, COALESCE(co.name,'') AS company "
           "FROM jobs j LEFT JOIN companies co ON co.id = j.company_id")
    if where:
        sql += " WHERE " + " AND ".join(where)
    sql += QUERY_ORDER
    return conn.execute(sql, params).fetchall()


def _export_csv(rows, path):
    import csv
    with open(path, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(EXPORT_COLS)
        for r in rows:
            w.writerow([r[c] for c in EXPORT_COLS])
    return path


def _md_table(rows):
    head = "| Tier | Company | Role | Location | Posted | Comp | Source | Why it fits | Screening risks |\n"
    head += "|---|---|---|---|---|---|---|---|---|\n"
    out = []
    for r in rows:
        comp = (f"{r['comp_min']//1000}-{r['comp_max']//1000}k" if r["comp_min"] and r["comp_max"]
                else (f"{r['comp_min']//1000}k+" if r["comp_min"] else ""))
        src = f"{r['verification_tag']} — {r['url'] or ''}"
        out.append(f"| {r['tier'] or '-'} | {r['company']} | {r['title']} | {r['location'] or ''} "
                   f"| {r['posting_date'] or '?'} | {comp} | {src} | {(r['fit_summary'] or '').replace(chr(10),' ')} "
                   f"| {(r['screening_risks'] or '').replace(chr(10),' ')} |")
    return head + "\n".join(out) + "\n"


def _export_md(rows, cand, path):
    matched = [r for r in rows if r["verification_tag"] != "wrong_location"]
    excluded = [r for r in rows if r["verification_tag"] == "wrong_location"]
    lines = [f"# Job Opportunities — {cand['name']}", ""]
    if cand["location_constraint"]:
        lines.append(f"**Location constraint:** {cand['location_constraint']}  ")
    lines.append(f"**Generated:** {today()} — {len(rows)} role(s), "
                 f"{len(matched)} location-matched, {len(excluded)} excluded.\n")
    for t, hdr in [(1, "🔴 Tier 1 — Apply Immediately"),
                   (2, "🟡 Tier 2 — Strong Fit, Pursue Actively"),
                   (3, "🟢 Tier 3 — Monitor / Opportunistic")]:
        tr = [r for r in matched if r["tier"] == t]
        if tr:
            lines += [f"## {hdr}", "", _md_table(tr), ""]
    untiered = [r for r in matched if not r["tier"]]
    if untiered:
        lines += ["## Untiered (aggregator / unverified)", "", _md_table(untiered), ""]
    if excluded:
        lines += ["## ❌ Excluded — Wrong Location", "", _md_table(excluded), ""]
    lines.append("> Re-verify each posting is still live AND in the right location before "
                 "applying — even verified-live roles can be pulled between report and application.")
    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return path


def _export_docx(rows, cand, path):
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        print("  (skipped .docx — install python-docx:  pip install python-docx)")
        return None
    doc = Document()
    doc.add_heading(f"Job Opportunities — {cand['name']}", level=0)
    if cand["location_constraint"]:
        doc.add_paragraph(f"Location constraint: {cand['location_constraint']}")
    matched = [r for r in rows if r["verification_tag"] != "wrong_location"]
    excluded = [r for r in rows if r["verification_tag"] == "wrong_location"]
    doc.add_paragraph(f"Generated {today()} — {len(rows)} role(s), {len(matched)} "
                      f"location-matched, {len(excluded)} excluded.")
    cols = ["Tier", "Company", "Role", "Location", "Posted", "Source", "Why it fits", "Risks"]

    def section(title, group):
        if not group:
            return
        doc.add_heading(title, level=1)
        tbl = doc.add_table(rows=1, cols=len(cols))
        tbl.style = "Light Grid Accent 1"
        for i, c in enumerate(cols):
            tbl.rows[0].cells[i].text = c
        for r in group:
            cells = tbl.add_row().cells
            vals = [str(r["tier"] or "-"), r["company"], r["title"], r["location"] or "",
                    r["posting_date"] or "?", f"{r['verification_tag']}\n{r['url'] or ''}",
                    r["fit_summary"] or "", r["screening_risks"] or ""]
            for i, v in enumerate(vals):
                cells[i].text = v
    for t, hdr in [(1, "Tier 1 — Apply Immediately"),
                   (2, "Tier 2 — Strong Fit"), (3, "Tier 3 — Monitor")]:
        section(hdr, [r for r in matched if r["tier"] == t])
    section("Untiered (aggregator / unverified)", [r for r in matched if not r["tier"]])
    section("Excluded — Wrong Location", excluded)
    doc.save(path)
    return path


# Native .xlsx writer (stdlib only — an xlsx is a zip of XML parts). Colors each row
# by status so the pipeline reads at a glance; freezes the header and adds an autofilter.
# Style indices below are referenced by the per-row s="" attribute in the sheet XML.
_XLSX_STATUS_XF = {  # status -> cellXfs index (see styles.xml below)
    "active": 2, "expired": 3, "new": 4, "applied": 5, "rejected": 6, "ignored": 7,
}
_XLSX_NUMERIC = {"id", "comp_min", "comp_max"}  # written as numbers so Excel sorts them


def _xlsx_col(n):  # 1-based column index -> letter(s)
    s = ""
    while n:
        n, rem = divmod(n - 1, 26)
        s = chr(65 + rem) + s
    return s


def _xlsx_esc(v):
    return (str(v).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            .replace('"', "&quot;"))


def _export_xlsx(rows, cand, path):
    import zipfile
    ncols = len(EXPORT_COLS)
    last = f"{_xlsx_col(ncols)}{len(rows) + 1}"

    def cell(col, ref, val, xf):
        if val is None or val == "":
            return f'<c r="{ref}" s="{xf}"/>'
        if col in _XLSX_NUMERIC and str(val).lstrip("-").isdigit():
            return f'<c r="{ref}" s="{xf}"><v>{val}</v></c>'
        return f'<c r="{ref}" s="{xf}" t="inlineStr"><is><t xml:space="preserve">{_xlsx_esc(val)}</t></is></c>'

    sd = ['<row r="1">']  # header row (style 1)
    sd += [cell(c, f"{_xlsx_col(i+1)}1", c, 1) for i, c in enumerate(EXPORT_COLS)]
    sd.append("</row>")
    for ri, r in enumerate(rows, start=2):
        xf = _XLSX_STATUS_XF.get(r["status"], 0)
        sd.append(f'<row r="{ri}">')
        sd += [cell(c, f"{_xlsx_col(i+1)}{ri}", r[c], xf) for i, c in enumerate(EXPORT_COLS)]
        sd.append("</row>")

    # helpful widths for the wide free-text columns (1-based positions in EXPORT_COLS)
    widths = {5: 22, 6: 40, 7: 26, 14: 46, 19: 50, 20: 44}
    cols_xml = "".join(f'<col min="{i}" max="{i}" width="{w}" customWidth="1"/>'
                       for i, w in widths.items())
    sheet = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">'
        '<sheetViews><sheetView workbookViewId="0">'
        '<pane ySplit="1" topLeftCell="A2" activePane="bottomLeft" state="frozen"/>'
        '</sheetView></sheetViews>'
        '<sheetFormatPr defaultRowHeight="15"/>'
        f'<cols>{cols_xml}</cols>'
        f'<sheetData>{"".join(sd)}</sheetData>'
        f'<autoFilter ref="A1:{last}"/>'
        '</worksheet>')

    def solid(rgb):
        return f'<fill><patternFill patternType="solid"><fgColor rgb="FF{rgb}"/></patternFill></fill>'
    styles = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<styleSheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">'
        '<fonts count="2">'
        '<font><sz val="11"/><name val="Calibri"/></font>'
        '<font><b/><color rgb="FFFFFFFF"/><sz val="11"/><name val="Calibri"/></font></fonts>'
        '<fills count="9">'
        '<fill><patternFill patternType="none"/></fill>'
        '<fill><patternFill patternType="gray125"/></fill>'
        + solid("305496") + solid("C6EFCE") + solid("FFC7CE") + solid("FFEB9C")
        + solid("BDD7EE") + solid("F8CBAD") + solid("D9D9D9") + '</fills>'
        '<borders count="1"><border/></borders>'
        '<cellStyleXfs count="1"><xf numFmtId="0" fontId="0" fillId="0" borderId="0"/></cellStyleXfs>'
        '<cellXfs count="8">'
        '<xf fontId="0" fillId="0" borderId="0"/>'
        '<xf fontId="1" fillId="2" borderId="0" applyFont="1" applyFill="1"/>'
        '<xf fillId="3" borderId="0" applyFill="1"/>'
        '<xf fillId="4" borderId="0" applyFill="1"/>'
        '<xf fillId="5" borderId="0" applyFill="1"/>'
        '<xf fillId="6" borderId="0" applyFill="1"/>'
        '<xf fillId="7" borderId="0" applyFill="1"/>'
        '<xf fillId="8" borderId="0" applyFill="1"/>'
        '</cellXfs></styleSheet>')

    parts = {
        "[Content_Types].xml":
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
            '<Default Extension="xml" ContentType="application/xml"/>'
            '<Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>'
            '<Override PartName="/xl/worksheets/sheet1.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>'
            '<Override PartName="/xl/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.styles+xml"/>'
            '</Types>',
        "_rels/.rels":
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>'
            '</Relationships>',
        "xl/workbook.xml":
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" '
            'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            '<sheets><sheet name="Pipeline" sheetId="1" r:id="rId1"/></sheets></workbook>',
        "xl/_rels/workbook.xml.rels":
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet1.xml"/>'
            '<Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>'
            '</Relationships>',
        "xl/styles.xml": styles,
        "xl/worksheets/sheet1.xml": sheet,
    }
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as z:
        for name, data in parts.items():
            z.writestr(name, data)
    return path


def cmd_export(args):
    conn = connect()
    cand = get_candidate(conn, args.candidate)
    rows = _export_rows(conn, args)
    conn.close()
    if not rows:
        print("(no matching jobs to export)")
        return
    os.makedirs(os.path.join(HERE, "exports"), exist_ok=True)
    base = args.out or os.path.join(HERE, "exports", f"{cand['slug']}_pipeline_{today()}")
    base = os.path.splitext(base)[0]  # strip any extension; we add per-format
    fmts = ["csv", "md", "xlsx", "docx"] if args.format == "all" else [args.format]
    written = []
    for fmt in fmts:
        path = f"{base}.{fmt}"
        if fmt == "csv":
            written.append(_export_csv(rows, path))
        elif fmt == "md":
            written.append(_export_md(rows, cand, path))
        elif fmt == "xlsx":
            written.append(_export_xlsx(rows, cand, path))
        elif fmt == "docx":
            res = _export_docx(rows, cand, path)
            if res:
                written.append(res)
    print(f"Exported {len(rows)} role(s):")
    for w in written:
        print(f"  {w}")
    if not args.all and not args.status:
        print("  (expired/rejected/ignored excluded — use --all for the full history)")


# ---------------------------------------------------------------------------
# company (list / rename / merge) — keeps the companies table clean when an
# employer is renamed, or when the same org was stored under two name strings
# (companies are keyed by exact name, so "NREL" and "NREL Inc." make two rows).
# ---------------------------------------------------------------------------
def _company_by_name(conn, name):
    return conn.execute("SELECT * FROM companies WHERE name = ?", (name,)).fetchone()


def cmd_company_list(args):
    conn = connect()
    rows = conn.execute(
        "SELECT co.id, co.name, co.ats_slug, co.careers_url, COUNT(j.id) AS jobs "
        "FROM companies co LEFT JOIN jobs j ON j.company_id = co.id "
        "GROUP BY co.id ORDER BY co.name"
    ).fetchall()
    conn.close()
    if not rows:
        print("(no companies)")
        return
    for r in rows:
        print(f"  id={r['id']:<4} jobs={r['jobs']:<3} {r['name']}  "
              f"[{r['ats_slug'] or '-'}]  {r['careers_url'] or ''}")


def cmd_company_add(args):
    """Register a target-list company WITHOUT needing a job for it (useful for small
    employers with no ATS feed that only post to LinkedIn — they still belong in the
    search scope so future sweeps don't skip them). Idempotent: if the name already
    exists, fills blanks / updates supplied fields instead of erroring or duplicating."""
    conn = connect()
    existing = _company_by_name(conn, args.name)
    cid = upsert_company(conn, {
        "company": args.name,
        "careers_url": args.careers_url,
        "ats_platform": args.ats_platform,
        "ats_slug": args.ats_slug,
        "warm_path": args.warm_path,
        "multi_region": args.multi_region,
    })
    if args.notes:  # guarded so an update without --notes never clobbers an existing note
        conn.execute("UPDATE companies SET notes=? WHERE id=?", (args.notes, cid))
    conn.commit()
    conn.close()
    print(f"{'Updated' if existing else 'Added'} company '{args.name}' (id={cid}).")


def cmd_company_rename(args):
    """Rename a company. If the target name already exists, MERGE the source into
    it (repoint its jobs, fold in any fields the target is missing, drop the source
    row) — so this both renames an employer and collapses duplicate rows."""
    conn = connect()
    src = _company_by_name(conn, args.from_name)
    if not src:
        sys.exit(f"No company named {args.from_name!r}. See: python jobsdb.py company list")
    extra = {}
    if args.careers_url:
        extra["careers_url"] = args.careers_url
    if args.ats_slug:
        extra["ats_slug"] = args.ats_slug

    dst = _company_by_name(conn, args.to_name) if args.to_name != args.from_name else None
    if dst and dst["id"] != src["id"]:
        n = conn.execute("SELECT COUNT(*) FROM jobs WHERE company_id = ?",
                         (src["id"],)).fetchone()[0]
        conn.execute("UPDATE jobs SET company_id = ? WHERE company_id = ?",
                     (dst["id"], src["id"]))
        sets = {f: src[f] for f in ("careers_url", "ats_platform", "ats_slug", "warm_path")
                if not dst[f] and src[f]}          # fill only the target's blanks
        if src["multi_region"]:
            sets["multi_region"] = 1               # sticky: only ever turn ON
        sets.update(extra)                         # explicit flags win
        if sets:
            conn.execute(f"UPDATE companies SET {', '.join(f'{k}=?' for k in sets)} WHERE id=?",
                         [*sets.values(), dst["id"]])
        conn.execute("DELETE FROM companies WHERE id = ?", (src["id"],))
        conn.commit()
        conn.close()
        print(f"Merged '{args.from_name}' into existing '{args.to_name}': "
              f"repointed {n} job(s), removed 1 duplicate row.")
        return

    sets = {"name": args.to_name, **extra}
    conn.execute(f"UPDATE companies SET {', '.join(f'{k}=?' for k in sets)} WHERE id=?",
                 [*sets.values(), src["id"]])
    conn.commit()
    conn.close()
    tail = " (careers_url/slug updated)" if extra else ""
    print(f"Renamed '{args.from_name}' -> '{args.to_name}'{tail}.")


def _company_job_count(conn, company_id):
    return conn.execute(
        "SELECT COUNT(*) FROM jobs WHERE company_id = ?", (company_id,)).fetchone()[0]


def cmd_company_show(args):
    """Inspect a company by exact name, or search with --like. The 'does this company
    exist / what do we know about it' lookup (company names are exact-keyed, so --like is
    how you find one when you don't know the stored string)."""
    if not args.name and not args.like:
        sys.exit("Provide a company NAME or --like <substring>.")
    if args.name and args.like:
        sys.exit("Provide either a company NAME or --like <substring>, not both.")
    conn = connect()
    if args.like:
        rows = conn.execute(
            "SELECT co.id, co.name, co.ats_slug, co.careers_url, co.verification_status, "
            "COUNT(j.id) AS jobs FROM companies co LEFT JOIN jobs j ON j.company_id = co.id "
            "WHERE co.name LIKE ? COLLATE NOCASE GROUP BY co.id ORDER BY co.name",
            (f"%{args.like}%",),
        ).fetchall()
        conn.close()
        if not rows:
            print("(no matches)")
            return
        for r in rows:
            print(f"  id={r['id']:<4} jobs={r['jobs']:<3} {r['name']}  "
                  f"[{r['ats_slug'] or '-'}]  <{r['verification_status'] or 'unchecked'}>  "
                  f"{r['careers_url'] or ''}")
        return
    co = _company_by_name(conn, args.name)
    if not co:
        conn.close()
        sys.exit(f"Company '{args.name}' not found. Try:  "
                 f"python jobsdb.py company show --like <part>")
    print(f"# {co['name']}  (id={co['id']})")
    for k in ("careers_url", "ats_platform", "ats_slug", "warm_path", "notes",
              "verification_status", "open_roles"):
        if co[k] not in (None, ""):
            print(f"  {k:<20} {co[k]}")
    if co["multi_region"]:
        print(f"  {'multi_region':<20} yes")
    print(f"  {'last_verified':<20} {verified_age(co['last_verified'])}")
    print(f"  {'jobs':<20} {_company_job_count(conn, co['id'])}")
    conn.close()


def cmd_company_verify(args):
    """Record a company-level verification outcome — the analog of `mark --verified` for a
    job. Create-or-update (the probe->verify path often meets a company not yet in the DB);
    --status is required so a bare call can't silently create an empty row. Identity/feed
    fields are sticky (filled via upsert_company, never clobbered); the verification state
    is always overwritten because it's the freshest truth."""
    if args.status not in COMPANY_VERIFY_STATUS:
        sys.exit(f"--status must be one of: {', '.join(sorted(COMPANY_VERIFY_STATUS))}")
    conn = connect()
    existed = _company_by_name(conn, args.name) is not None
    cid = upsert_company(conn, {
        "company": args.name,
        "ats_platform": args.ats_platform,
        "ats_slug": args.ats_slug,
        "careers_url": args.careers_url,
    })
    updates = {"verification_status": args.status, "last_verified": args.date or today()}
    if args.open_roles is not None:
        updates["open_roles"] = args.open_roles
    if args.note:  # append a dated line, like `mark --note`, rather than overwrite
        row = conn.execute("SELECT notes FROM companies WHERE id=?", (cid,)).fetchone()
        prev = (row["notes"] + "\n") if row["notes"] else ""
        updates["notes"] = prev + f"[{today()}] {args.note}"
    conn.execute(f"UPDATE companies SET {', '.join(f'{k}=?' for k in updates)} WHERE id=?",
                 [*updates.values(), cid])
    conn.commit()
    conn.close()
    roles = f", open_roles={args.open_roles}" if args.open_roles is not None else ""
    print(f"{'Updated' if existed else 'Created'} company '{args.name}' (id={cid}): "
          f"status={args.status}, last_verified={updates['last_verified']}{roles}.")


# ---------------------------------------------------------------------------
# arg parsing
# ---------------------------------------------------------------------------
def build_parser():
    p = argparse.ArgumentParser(prog="jobsdb.py", description="Local job-pipeline database.")
    sub = p.add_subparsers(dest="cmd", required=True)

    sp = sub.add_parser("init", help="Create the database from schema.sql")
    sp.add_argument("--force", action="store_true", help="Drop and recreate (DESTROYS data)")
    sp.set_defaults(func=cmd_init)

    cand = sub.add_parser("candidate", help="Manage candidates")
    csub = cand.add_subparsers(dest="action", required=True)
    ca = csub.add_parser("add", help="Add or update a candidate")
    ca.add_argument("--resume")
    ca.add_argument("--slug")
    ca.add_argument("--field", action="append", help="key=value (repeatable)")
    ca.set_defaults(func=cmd_candidate_add)
    cl = csub.add_parser("list"); cl.set_defaults(func=cmd_candidate_list)
    cs = csub.add_parser("show"); cs.add_argument("--slug", required=True)
    cs.set_defaults(func=cmd_candidate_show)

    cat = sub.add_parser("category", help="Manage candidate categories")
    catsub = cat.add_subparsers(dest="action", required=True)
    cset = catsub.add_parser("set", help="Replace a candidate's ranked categories")
    cset.add_argument("--candidate", required=True)
    cset.add_argument("--json", required=True, help="JSON file: [{rank,label,keywords}]")
    cset.set_defaults(func=cmd_category_set)

    co = sub.add_parser("company",
                        help="Manage companies (list / show / add / verify / rename / merge)")
    cosub = co.add_subparsers(dest="action", required=True)
    cll = cosub.add_parser("list", help="List companies with job counts")
    cll.set_defaults(func=cmd_company_list)
    csh = cosub.add_parser("show", help="Inspect a company by name, or search with --like")
    csh.add_argument("name", nargs="?", help="Exact company name to inspect")
    csh.add_argument("--like", help="Case-insensitive substring search instead of exact name")
    csh.set_defaults(func=cmd_company_show)
    cvf = cosub.add_parser("verify",
                           help="Record a company verification outcome (analog of job mark --verified)")
    cvf.add_argument("name")
    cvf.add_argument("--status", required=True, choices=sorted(COMPANY_VERIFY_STATUS))
    cvf.add_argument("--date", help="ISO date of the check (default: today)")
    cvf.add_argument("--open-roles", dest="open_roles", type=int,
                     help="Open-role count from the probe")
    cvf.add_argument("--ats-platform", dest="ats_platform")
    cvf.add_argument("--ats-slug", dest="ats_slug")
    cvf.add_argument("--careers-url", dest="careers_url")
    cvf.add_argument("--note", help="Append a dated note line")
    cvf.set_defaults(func=cmd_company_verify)
    cad = cosub.add_parser("add", help="Register a target-list company (no job needed)")
    cad.add_argument("--name", required=True)
    cad.add_argument("--careers-url", dest="careers_url")
    cad.add_argument("--ats-platform", dest="ats_platform")
    cad.add_argument("--ats-slug", dest="ats_slug")
    cad.add_argument("--multi-region", dest="multi_region", action="store_true")
    cad.add_argument("--warm-path", dest="warm_path")
    cad.add_argument("--notes")
    cad.set_defaults(func=cmd_company_add)
    crn = cosub.add_parser("rename",
                           help="Rename a company; merges into the target if it already exists")
    crn.add_argument("--from", dest="from_name", required=True)
    crn.add_argument("--to", dest="to_name", required=True)
    crn.add_argument("--careers-url", dest="careers_url", help="Also update careers_url")
    crn.add_argument("--ats-slug", dest="ats_slug", help="Also update ats_slug")
    crn.set_defaults(func=cmd_company_rename)

    ub = sub.add_parser("upsert-batch",
                        help="Insert/update jobs from a scan batch (job_scans/*.json)")
    ub.add_argument("file")
    ub.set_defaults(func=cmd_upsert_batch)

    q = sub.add_parser("query", help="Query the pipeline")
    q.add_argument("--candidate")
    q.add_argument("--category")
    q.add_argument("--tier", type=int, choices=[1, 2, 3])
    q.add_argument("--status", choices=sorted(VALID_STATUS))
    q.add_argument("--verification", choices=sorted(VALID_TAGS))
    q.add_argument("--location-match", dest="location_match", choices=["yes", "no"])
    q.add_argument("--since")
    q.add_argument("--limit", type=int)
    q.add_argument("--all", action="store_true",
                   help="Include expired/rejected/ignored (hidden by default)")
    q.add_argument("--format", choices=["table", "json"], default="table")
    q.set_defaults(func=cmd_query)

    st = sub.add_parser("stats", help="Pipeline breakdown")
    st.add_argument("--candidate", required=True)
    st.set_defaults(func=cmd_stats)

    rv = sub.add_parser("reverify", help="List stale jobs needing re-verification")
    rvsub = rv.add_subparsers(dest="action", required=True)
    rvl = rvsub.add_parser("list", help="Emit stale (new/active) jobs to re-check")
    rvl.add_argument("--candidate", required=True)
    rvl.add_argument("--stale-days", type=int, default=2,
                     help="Re-check jobs whose last_verified is older than N days "
                          "(default 2). Tier 1/2 are always re-checked unless verified today.")
    rvl.add_argument("--format", choices=["table", "json"], default="table")
    rvl.set_defaults(func=cmd_reverify)

    mk = sub.add_parser("mark", help="Update one job (status / verified / paths / notes)")
    mk.add_argument("job_id", type=int)
    mk.add_argument("--status", choices=sorted(VALID_STATUS))
    mk.add_argument("--verified", action="store_true", help="Refresh last_verified to today")
    mk.add_argument("--resume", help="Path to tailored resume .docx")
    mk.add_argument("--cover", help="Path to cover letter .docx")
    mk.add_argument("--applied-date", dest="applied_date", help="ISO date")
    mk.add_argument("--note", help="Append a timestamped note")
    mk.set_defaults(func=cmd_mark)

    ex = sub.add_parser("export", help="Export the pipeline (csv/md/xlsx/docx) FROM the DB")
    ex.add_argument("--candidate", required=True)
    ex.add_argument("--format", choices=["csv", "md", "xlsx", "docx", "all"], default="all")
    ex.add_argument("--out", help="Output path base (extension added per format)")
    ex.add_argument("--category"); ex.add_argument("--tier", type=int, choices=[1, 2, 3])
    ex.add_argument("--status", choices=sorted(VALID_STATUS))
    ex.add_argument("--verification", choices=sorted(VALID_TAGS))
    ex.add_argument("--location-match", dest="location_match", choices=["yes", "no"])
    ex.add_argument("--since")
    ex.add_argument("--all", action="store_true",
                    help="Include expired/rejected/ignored (hidden by default)")
    ex.set_defaults(func=cmd_export)

    return p


def main():
    args = build_parser().parse_args()
    args.func(args)


if __name__ == "__main__":
    main()
